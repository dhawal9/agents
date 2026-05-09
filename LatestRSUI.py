import asyncio
import json
import re
from datetime import date
from docx import Document

import streamlit as st
import plotly.graph_objects as go
from pypdf import PdfReader

from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.teams import RoundRobinGroupChat
from autogen_ext.models.openai import OpenAIChatCompletionClient


model_client = OpenAIChatCompletionClient(
    model="gpt-4o-mini",
    api_key=st.secrets["OPENAI_API_KEY"],
    model_info={
        "family": "gpt-4",
        "vision": True,
        "function_calling": True,
        "json_output": True,
        "structured_output": True,
    }
)


def extract_text_from_pdf(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    text_parts = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    return "\n".join(text_parts).strip()


def extract_text_from_docx(uploaded_file) -> str:
    doc_bytes = uploaded_file.read()
    document = Document(io.BytesIO(doc_bytes))

    text_parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Extract text from tables also
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts).strip()


def extract_resume_text(uploaded_file) -> str:
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    elif file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    elif file_name.endswith(".doc"):
        raise ValueError(
            "Legacy .doc files are not supported. Please convert the resume to .docx or PDF."
        )

    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")


MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}


def parse_month_year(value):
    value = value.strip().lower()

    if value in ["present", "current", "till date", "to date"]:
        today = date.today()
        return today.year, today.month

    match = re.search(r"([a-zA-Z]+)\s+(\d{4})", value)
    if match:
        month_text = match.group(1).lower()
        year = int(match.group(2))
        if month_text in MONTHS:
            return year, MONTHS[month_text]

    match = re.search(r"\b(\d{4})\b", value)
    if match:
        return int(match.group(1)), 1

    return None


def to_month_index(year_month):
    year, month = year_month
    return year * 12 + month


def month_index_to_label(month_index):
    year = month_index // 12
    month = month_index % 12

    if month == 0:
        year -= 1
        month = 12

    month_names = [
        "", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
        "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"
    ]

    return f"{month_names[month]} {year}"


def merge_periods(periods):
    if not periods:
        return []

    periods = sorted(periods)
    merged = [periods[0]]

    for start, end in periods[1:]:
        last_start, last_end = merged[-1]

        if start <= last_end + 1:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))

    return merged


def extract_experience_section(resume_text):
    text = resume_text

    start_keywords = [
        "professional experience",
        "work experience",
        "employment history",
        "experience",
        "career history"
    ]

    end_keywords = [
        "education",
        "academic",
        "skills",
        "technical skills",
        "projects",
        "certifications",
        "achievements",
        "summary"
    ]

    lower_text = text.lower()

    start_index = -1
    for keyword in start_keywords:
        idx = lower_text.find(keyword)
        if idx != -1:
            start_index = idx
            break

    if start_index == -1:
        return text

    end_index = len(text)

    for keyword in end_keywords:
        idx = lower_text.find(keyword, start_index + 20)
        if idx != -1:
            end_index = min(end_index, idx)

    return text[start_index:end_index]


def calculate_experience_from_resume(resume_text):
    experience_text = extract_experience_section(resume_text)

    date_pattern = re.compile(
        r"((?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{4})"
        r"\s*(?:-|–|—|to)\s*"
        r"((?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{4}|Present|Current|Till Date|To Date)",
        re.IGNORECASE
    )

    periods = []
    raw_periods = []

    for match in date_pattern.finditer(experience_text):
        start_text = match.group(1)
        end_text = match.group(2)

        start = parse_month_year(start_text)
        end = parse_month_year(end_text)

        if not start or not end:
            continue

        start_index = to_month_index(start)
        end_index = to_month_index(end)

        if end_index < start_index:
            continue

        months = end_index - start_index + 1

        # Ignore unusually long single ranges, likely education or non-employment
        if months > 180:
            continue

        periods.append((start_index, end_index))
        raw_periods.append({
            "raw_text": match.group(0),
            "start": month_index_to_label(start_index),
            "end": month_index_to_label(end_index),
            "duration_months": months
        })

    merged_periods = merge_periods(periods)
    total_months = sum(end - start + 1 for start, end in merged_periods)

    merged_period_labels = [
        {
            "start": month_index_to_label(start),
            "end": month_index_to_label(end),
            "duration_months": end - start + 1
        }
        for start, end in merged_periods
    ]

    return {
        "years_of_experience": round(total_months / 12, 1) if total_months > 0 else None,
        "total_months": total_months,
        "raw_periods_found": raw_periods,
        "merged_periods": merged_period_labels,
        "calculation_method": "Python deterministic calculation from professional experience section only",
        "note": "Only month-year employment ranges from the experience section are counted. Education/projects/certifications are excluded."
    }


def extract_json_from_text(content):
    if not content:
        return None

    try:
        return json.loads(content)
    except Exception:
        pass

    start = content.find("{")
    end = content.rfind("}")

    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(content[start:end + 1])
        except Exception:
            return None

    return None


def extract_agent_outputs(messages):
    jd_data = {}
    parser_data = {}
    matcher_data = {}
    screening_data = {}
    summary = ""

    for msg in messages:
        content = getattr(msg, "content", "")

        if not content:
            continue

        data = extract_json_from_text(content)

        if isinstance(data, dict):
            if "required_skills" in data and "preferred_skills" in data:
                jd_data = data
            elif "years_of_experience" in data and "skills" in data:
                parser_data = data
            elif "overall_match_score" in data:
                matcher_data = data
            elif "decision" in data:
                screening_data = data
        else:
            clean_content = content.replace("TERMINATE", "").strip()
            if clean_content:
                summary = clean_content

    return jd_data, parser_data, matcher_data, screening_data, summary


def safe_number(value, default=0):
    try:
        if value is None:
            return default
        return float(value)
    except Exception:
        return default


def render_gauge(title, value):
    value = safe_number(value)

    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=value,
        title={"text": title},
        gauge={
            "axis": {"range": [0, 100]},
            "bar": {"color": "royalblue"},
            "steps": [
                {"range": [0, 40], "color": "#ffcccc"},
                {"range": [40, 70], "color": "#fff3cd"},
                {"range": [70, 100], "color": "#d4edda"},
            ],
        }
    ))

    fig.update_layout(height=300, margin=dict(l=20, r=20, t=50, b=20))
    st.plotly_chart(fig, use_container_width=True)


def render_donut(labels, values, title):
    fig = go.Figure(data=[
        go.Pie(
            labels=labels,
            values=values,
            hole=0.55
        )
    ])

    fig.update_layout(
        title=title,
        height=350,
        margin=dict(l=20, r=20, t=50, b=20),
        annotations=[
            dict(
                text=title,
                x=0.5,
                y=0.5,
                font_size=16,
                showarrow=False
            )
        ]
    )

    st.plotly_chart(fig, use_container_width=True)


def render_score_breakdown(score_breakdown):
    if not score_breakdown:
        st.info("No detailed score breakdown available.")
        return

    categories = []
    scores = []

    for item in score_breakdown:
        category = item.get("category", "Unknown")
        score = safe_number(item.get("score", 0))
        max_score = safe_number(item.get("max_score", 100))

        percentage = 0
        if max_score > 0:
            percentage = round((score / max_score) * 100, 2)

        categories.append(category)
        scores.append(percentage)

    fig = go.Figure(go.Bar(
        x=scores,
        y=categories,
        orientation="h",
        text=[f"{s}%" for s in scores],
        textposition="auto"
    ))

    fig.update_layout(
        title="Detailed Score Breakdown",
        xaxis_title="Score %",
        yaxis_title="Category",
        xaxis=dict(range=[0, 100]),
        height=400,
        margin=dict(l=20, r=20, t=50, b=20)
    )

    st.plotly_chart(fig, use_container_width=True)


def create_agents(model_client):

    resume_parser = AssistantAgent(
    name="ResumeParserAgent",
    model_client=model_client,
    system_message="""
You are a Resume Parsing Specialist.

Extract structured information from the resume.

Important:
- Do not infer beyond the resume content.
- Do not calculate total experience yourself.
- If Python-calculated experience is provided in the input, use it as the source of truth.
- If explicit experience is mentioned in the resume, capture it separately.
- Output JSON only.

Output MUST be valid JSON only:
{
  "years_of_experience": null,
  "python_calculated_experience_years": null,
  "explicit_experience_mentioned": null,
  "experience_evidence": "",
  "skills": [
    {
      "skill": "",
      "evidence": ""
    }
  ],
  "technologies": [
    {
      "technology": "",
      "evidence": ""
    }
  ],
  "cloud_experience": [
    {
      "platform": "",
      "evidence": ""
    }
  ],
  "projects": [
    {
      "name": "",
      "description": "",
      "technologies_used": [],
      "evidence": ""
    }
  ],
  "education": "",
  "education_evidence": "",
  "employment_history": [
    {
      "company": "",
      "role": "",
      "start_date": "",
      "end_date": "",
      "evidence": ""
    }
  ],
  "summary": ""
}

Experience Rules:
- years_of_experience must be equal to python_calculated_experience_years if provided.
- Do not recalculate years_of_experience using the LLM.
- If the resume explicitly says something like "4+ years experience", store it in explicit_experience_mentioned.
- Put the exact phrase in experience_evidence.
- If neither Python-calculated nor explicit experience is available, use null.

Skill Extraction Rules:
- Extract only skills explicitly present in the resume.
- Every skill must include evidence from the resume.
- Do not add related skills.
- Do not infer Spring Boot from Java.
- Do not infer Java from Spring Boot.
- Do not infer SQL from Oracle/MySQL/PostgreSQL unless SQL/query writing is explicitly mentioned.
- Do not infer AWS from cloud.
- Do not infer Docker/Kubernetes from DevOps unless explicitly mentioned.
- Keep skills and technologies factual.

Evidence Rules:
- Evidence must be short text copied or closely quoted from the resume.
- Do not use JD content as resume evidence.
- If evidence is unavailable, do not include that skill.
- Avoid duplicate skills.

Summary Rules:
- Summary should be factual and based only on parsed resume content.
- Do not include recommendation or match score here.
"""
)

    jd_parser = AssistantAgent(
    name="JDParserAgent",
    model_client=model_client,
    system_message="""
You are a Job Description Parsing Specialist.

Extract structured hiring requirements from the job description.

Your output will be used by a resume matcher, so be strict and evidence-based.

Output MUST be valid JSON only:
{
  "job_title": "",
  "required_skills": [
    {
      "skill": "",
      "evidence": ""
    }
  ],
  "preferred_skills": [
    {
      "skill": "",
      "evidence": ""
    }
  ],
  "required_experience_years": null,
  "experience_evidence": "",
  "required_education": "",
  "education_evidence": "",
  "responsibilities": [],
  "domain": "",
  "must_have_requirements": [],
  "nice_to_have_requirements": [],
  "raw_jd_skill_terms": []
}

Rules:
- Output JSON only.
- Do not hallucinate requirements.
- If something is not explicitly mentioned, return null, "", or [].
- Extract only skills explicitly mentioned in the JD.
- Do not add related skills.
- Do not infer frameworks from programming languages.
- Do not infer programming languages from frameworks.
- Separate mandatory and optional requirements carefully.

Required skills:
- Add a skill to required_skills only if the JD clearly marks it as mandatory.
- Indicators include: required, must have, mandatory, should have, strong experience in, hands-on experience in, responsible for, expected to work on.

Preferred skills:
- Add a skill to preferred_skills only if the JD clearly marks it as optional or good-to-have.
- Indicators include: preferred, nice to have, good to have, advantage, plus, optional, exposure to.

Evidence rules:
- Every required_skills item must include exact evidence from the JD.
- Every preferred_skills item must include exact evidence from the JD.
- Evidence should be a short phrase copied from the JD.
- Do not include resume evidence here.

Experience rules:
- Extract required_experience_years only from explicit JD wording.
- Examples:
  - "4+ years experience" -> 4
  - "minimum 5 years" -> 5
  - "3-6 years" -> 3
  - "at least 7 years" -> 7
- If experience is not mentioned, use null.
- Put the exact JD phrase in experience_evidence.

Skill normalization rules:
- Keep the skill name concise.
- Example: "hands-on experience in Java and Spring Boot" should become:
  - Java
  - Spring Boot
- Do not expand "Java" into Spring Boot, Hibernate, Microservices, or REST API unless those are explicitly mentioned.
- Do not expand "cloud" into AWS/Azure/GCP unless explicitly mentioned.
- Do not expand "database" into SQL/Oracle/MySQL unless explicitly mentioned.
- Do not expand "CI/CD" into Jenkins/Docker/Kubernetes unless explicitly mentioned.

raw_jd_skill_terms:
- Include a flat list of all exact skill terms found in the JD.
- This list should contain required and preferred skill terms only.
- This list will be used as the allowed matching universe.
"""
)

    

    jd_matcher = AssistantAgent(
    name="EvidenceBasedMatcherAgent",
    model_client=model_client,
    system_message="""
You are an Evidence-Based Resume Matching Engine.

Compare the parsed JD and parsed resume strictly against the skills explicitly mentioned in the JD.

Output MUST be valid JSON only:
{
  "overall_match_score": number,
  "required_skill_evaluation": [
    {
      "skill": "",
      "status": "matched | missing",
      "evidence": "",
      "reason": ""
    }
  ],
  "preferred_skill_evaluation": [
    {
      "skill": "",
      "status": "matched | missing",
      "evidence": "",
      "reason": ""
    }
  ],
  "matched_skills": [
    {
      "skill": "",
      "jd_requirement_type": "required | preferred",
      "evidence": ""
    }
  ],
  "missing_required_skills": [],
  "missing_preferred_skills": [],
  "resume_only_skills_ignored_for_scoring": [],
  "score_breakdown": [
    {
      "category": "Technical Skills",
      "score": number,
      "max_score": 40,
      "reason": ""
    },
    {
      "category": "Experience",
      "score": number,
      "max_score": 25,
      "reason": ""
    },
    {
      "category": "Domain Relevance",
      "score": number,
      "max_score": 15,
      "reason": ""
    },
    {
      "category": "Projects",
      "score": number,
      "max_score": 10,
      "reason": ""
    },
    {
      "category": "Education",
      "score": number,
      "max_score": 5,
      "reason": ""
    },
    {
      "category": "Resume Quality",
      "score": number,
      "max_score": 5,
      "reason": ""
    }
  ],
  "strengths": [],
  "concerns": [],
  "recommendation_reason": ""
}

Mandatory Skill Audit Rules:
- You must evaluate every required skill from the JD.
- You must evaluate every preferred skill from the JD.
- Every required JD skill must appear exactly once in required_skill_evaluation.
- Every preferred JD skill must appear exactly once in preferred_skill_evaluation.
- Do not skip preferred skills.
- Do not leave preferred_skill_evaluation empty unless the JD has no preferred skills.

Matching Rules:
- A skill can be marked as matched only if there is clear resume evidence.
- Evidence must come only from the resume.
- If evidence is missing, weak, indirect, or only assumed, status must be "missing".
- Do not infer related skills.
- Do not infer Java from Spring Boot.
- Do not infer Spring Boot from Java.
- Do not infer SQL from Oracle/MySQL/PostgreSQL unless SQL/query writing is explicitly mentioned.
- Do not infer AWS/Azure/GCP from the word cloud.
- Do not infer Docker/Kubernetes from DevOps.
- Do not match resume-only skills that are not present in the JD.

Consistency Rules:
- missing_required_skills must contain exactly those skills from required_skill_evaluation where status = "missing".
- missing_preferred_skills must contain exactly those skills from preferred_skill_evaluation where status = "missing".
- matched_skills must contain exactly those skills from required_skill_evaluation and preferred_skill_evaluation where status = "matched".
- Do not put a skill in matched_skills unless it appears in required_skill_evaluation or preferred_skill_evaluation.
- Do not put resume-only skills in matched_skills.
- If a preferred JD skill has no evidence, it must appear in missing_preferred_skills.
- Do not leave missing_preferred_skills empty unless every preferred skill has status "matched".

Resume-only Skill Rule:
- If the resume has skills not mentioned in the JD, list them under resume_only_skills_ignored_for_scoring.
- Resume-only skills must not increase score.
- Resume-only skills must not be added to matched_skills.

Critical Experience Rule:
- Use Python-calculated experience as the source of truth if provided.
- Do not recalculate candidate experience.
- Do not reduce total experience because one technology has fewer years.
- If required experience is satisfied by Python-calculated experience, do not reject due to experience.

Self-Check Before Final JSON:
Before producing the final JSON, verify internally:
1. Count of required_skill_evaluation equals count of JD required_skills.
2. Count of preferred_skill_evaluation equals count of JD preferred_skills.
3. Every missing preferred skill appears in missing_preferred_skills.
4. No resume-only skill appears in matched_skills.
5. overall_match_score equals the sum of all score_breakdown scores.

Scoring Rules:
- Total max score is 100.
- overall_match_score must equal the sum of all score_breakdown scores.
- Missing required skills should affect the score heavily.
- Missing preferred skills should reduce score, but by at most 10 points total.
- If all required JD skills are matched and required experience is satisfied, Technical Skills should usually be 36-40 out of 40.
- If required experience is satisfied, Experience should usually be 23-25 out of 25.
- If no required skills are missing and experience is satisfied, overall_match_score should usually be >= 85.
- If all required and preferred JD skills are matched, overall_match_score should usually be >= 90.
- Reserve scores below 80 for candidates with clear required-skill or experience gaps.

Output JSON only.
"""
)

    screening_agent = AssistantAgent(
        name="ScreeningAgent",
        model_client=model_client,
        system_message="""
You are a Hiring Decision Engine.

Use the evidence-based matching output to make a decision.

Output MUST be valid JSON only:
{
  "decision": "HIRE" | "REJECT" 
  ,
  "confidence": number,
  "reasoning": ""
}

Critical Experience Rule:
- Use Python-calculated experience as the source of truth.
- Do not reject a candidate for experience if Python-calculated experience satisfies the JD requirement.
- Major missing mandatory skills can still lead to REJECT.

Rules:
- HIRE usually requires score >= 75 and no major required skill gap.
- REJECT is for weak matches or major missing mandatory requirements.
- Be objective and bias-free.
- Output JSON only.

Confidence Calibration Rules:
- Confidence reflects how certain you are about the decision, not the candidate quality.
- If evidence is clear and score is high, confidence should be 90+.
- If all required skills and required experience are matched, decision should usually be HIRE.
- Do not set confidence to 80 by default.
"""
    )

    summary_agent = AssistantAgent(
        name="SummaryAgent",
        model_client=model_client,
        system_message="""
You are a Recruiter-Facing Summary Generator.

Create a short professional summary based on all previous outputs.

Format:
### Candidate Overview
### Key Strengths
### Skill Gaps
### Final Recommendation

Rules:
- Keep it concise.
- Use professional HR language.
- No JSON output.
- Mention that total experience is Python-calculated if experience is discussed.

IMPORTANT:
After generating the final summary, respond with:
TERMINATE
"""
    )

    return RoundRobinGroupChat(
        participants=[
            jd_parser,
            resume_parser,
            jd_matcher,
            screening_agent,
            summary_agent
        ],
        max_turns=5
    )


async def run_agents(groupchat, jd, resume_text, experience_result):
    result = []

    async for message in groupchat.run_stream(
        task=f"""
You must analyze the following job description and resume.

Important:
The following Python-calculated experience is the source of truth for total professional experience.
Do not recalculate total years of experience using the LLM.

Python-calculated experience:
{json.dumps(experience_result, indent=2)}

Job Description:
{jd}

Resume:
{resume_text}
"""
    ):
        result.append(message)

    return result


def run_agents_sync(groupchat, jd, resume_text, experience_result):
    return asyncio.run(run_agents(groupchat, jd, resume_text, experience_result))


st.set_page_config(
    page_title="AI Resume Screener",
    layout="wide"
)

st.title("📄 AI Resume Screening Dashboard")
st.caption("Evidence-based resume screening with deterministic Python experience calculation.")

jd = st.text_area(
    "Enter Job Description",
    height=220,
    placeholder="Example: Java developer with Spring Boot, Kafka, SQL, cloud experience..."
)

uploaded_file = st.file_uploader(
    "Upload Resume",
    type=["pdf", "docx"]
)

if st.button("Analyze Candidate", type="primary"):

    if not jd or not uploaded_file:
        st.warning("Please provide both Job Description and Resume PDF.")

    else:
        with st.spinner("Analyzing candidate profile..."):

            resume_text = extract_text_from_pdf(uploaded_file)

            if not resume_text:
                st.error("Could not extract text from the PDF. Please upload a text-based resume PDF.")
                st.stop()

            experience_result = calculate_experience_from_resume(resume_text)

            groupchat = create_agents(model_client)

            messages = run_agents_sync(
                groupchat,
                jd,
                resume_text,
                experience_result
            )

            jd_data, parser_data, matcher_data, screening_data, summary = extract_agent_outputs(messages)

        st.success("Analysis Complete!")

        match_score = safe_number(matcher_data.get("overall_match_score", 0))
        confidence = safe_number(screening_data.get("confidence", 0))
        decision = screening_data.get("decision", "N/A")

        python_experience = experience_result.get("years_of_experience")
        experience = python_experience if python_experience is not None else "N/A"

        candidate_name = parser_data.get("candidate_name") or "Candidate"

        matched_skills = matcher_data.get("matched_skills", [])
        missing_required = matcher_data.get("missing_required_skills", [])
        missing_preferred = matcher_data.get("missing_preferred_skills", [])
        score_breakdown = matcher_data.get("score_breakdown", [])

        st.subheader("📊 Screening Snapshot")

        col1, col2, col3, col4 = st.columns(4)

        col1.metric("Candidate", candidate_name)
        col2.metric("Match Score", f"{int(match_score)}%")
        col3.metric("Decision", decision)
        col4.metric("Confidence", f"{int(confidence)}%")

        col5, col6, col7, col8 = st.columns(4)

        col5.metric("Python Calculated Experience", f"{experience} yrs")
        col6.metric("Matched Skills", len(matched_skills))
        col7.metric("Missing Required", len(missing_required))
        col8.metric("Missing Preferred", len(missing_preferred))

        st.divider()

        tab1, tab2, tab3 = st.tabs([
            "🧾 Summary",
            "🧠 Skills & Evidence",
            "📈 Score Breakdown"
        ])

        with tab1:
            st.subheader("Recruiter Summary")

            if summary:
                st.markdown(summary)
            else:
                st.info("No final summary was generated.")

            st.subheader("Experience Calculation")

            st.write("**Experience Source:** Python deterministic calculation")
            st.write("**Calculated Years:**", experience_result.get("years_of_experience"))
            st.write("**Total Months:**", experience_result.get("total_months"))

            with st.expander("View experience periods used for calculation"):
                st.json({
                    "raw_periods_found": experience_result.get("raw_periods_found", []),
                    "merged_periods": experience_result.get("merged_periods", []),
                    "note": experience_result.get("note")
                })

            st.subheader("Parsed Job Requirement Summary")

            col1, col2 = st.columns(2)

            with col1:
                st.write("**Job Title:**", jd_data.get("job_title", "Not available"))
                st.write("**Domain:**", jd_data.get("domain", "Not available"))
                st.write("**Required Experience:**", jd_data.get("required_experience_years", "Not available"))

            with col2:
                st.write("**Required Education:**", jd_data.get("required_education", "Not available"))
                st.write("**Candidate Experience:**", f"{experience} yrs")

            st.subheader("Hiring Reasoning")
            st.write(screening_data.get("reasoning", "No reasoning available."))

            st.subheader("Matcher Recommendation Reason")
            st.write(matcher_data.get("recommendation_reason", "No recommendation reason available."))

        with tab2:
            col1, col2 = st.columns(2)

            with col1:
                st.subheader("✅ Matched Skills With Evidence")

                if matched_skills:
                    for item in matched_skills:
                        if isinstance(item, dict):
                            skill = item.get("skill", "Unknown skill")
                            evidence = item.get("evidence", "No evidence provided.")
                            with st.expander(f"✅ {skill}"):
                                st.write(evidence)
                        else:
                            st.success(str(item))
                else:
                    st.info("No matched skills found.")

            with col2:
                st.subheader("❌ Missing Skills")

                st.write("**Missing Required Skills**")
                if missing_required:
                    for skill in missing_required:
                        st.error(skill)
                else:
                    st.success("No required skills missing.")

                st.write("**Missing Preferred Skills**")
                if missing_preferred:
                    for skill in missing_preferred:
                        st.warning(skill)
                else:
                    st.success("No preferred skills missing.")

            st.divider()

            st.subheader("Skill Match Distribution")

            matched_count = len(matched_skills)
            missing_required_count = len(missing_required)
            missing_preferred_count = len(missing_preferred)

            if matched_count + missing_required_count + missing_preferred_count > 0:
                render_donut(
                    labels=["Matched", "Missing Required", "Missing Preferred"],
                    values=[matched_count, missing_required_count, missing_preferred_count],
                    title="Skills"
                )
            else:
                st.info("Skill distribution could not be calculated.")

            col1, col2 = st.columns(2)

            with col1:
                st.subheader("💪 Strengths")
                strengths = matcher_data.get("strengths", [])

                if strengths:
                    for strength in strengths:
                        st.write(f"✔️ {strength}")
                else:
                    st.info("No strengths found.")

            with col2:
                st.subheader("⚠️ Concerns")
                concerns = matcher_data.get("concerns", [])

                if concerns:
                    for concern in concerns:
                        st.write(f"⚠️ {concern}")
                else:
                    st.info("No concerns found.")

        with tab3:
            st.subheader("Visual Evaluation")

            col1, col2 = st.columns(2)

            with col1:
                render_gauge("Overall Match Score", match_score)

            with col2:
                render_gauge("Decision Confidence", confidence)

            st.divider()

            render_score_breakdown(score_breakdown)

            st.subheader("Score Explanation")

            if score_breakdown:
                for item in score_breakdown:
                    category = item.get("category", "Unknown")
                    score = item.get("score", 0)
                    max_score = item.get("max_score", 0)
                    reason = item.get("reason", "No reason available.")

                    with st.expander(f"{category}: {score}/{max_score}"):
                        st.write(reason)
            else:
                st.info("No score explanation available.")